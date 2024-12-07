
import docx.enum
from .stuff import *

class oilWindow():
    def showOil(self):
        try:
            self.destroyFrame(self.mainMenuOilFrame)
            self.destroyFrame(self.oilFrame)
        except:
            pass
        
        self.oilFrame = QFrame(self.mainFrame)
        
        self.closeButtonOilFrame = QPushButton(self.oilFrame)
        self.closeButtonOilFrame.setStyleSheet("QPushButton {border:2px solid black;font-size:12px;qproperty-icon: url('assests/close.png');qproperty-iconSize: 26px}QPushButton:hover {background-color:#c8c8c8;}")        
        self.closeButtonOilFrame.clicked.connect(lambda x, frame=self.oilFrame:self.destroyFrame(frame))

        self.comboSearchOilBox = QComboBox(self.oilFrame)
        self.comboSearchOilBox.setGeometry(470,30,150,20)
        self.comboSearchOilBox.addItems(["الكل", "كود الصنف"])
        self.comboSearchOilBox.activated.connect(self.addSearchOilEntries)

        self.comboOilStatusBox = QComboBox(self.oilFrame)
        self.comboOilStatusBox.setGeometry(100,30,150,20)
        self.comboOilStatusBox.addItems(["الكل", "جديد","تالف","تم البيع"])
        self.comboOilStatusBox.activated.connect(self.loadOil)

        self.oilFrame.setStyleSheet("background-color:white")

        self.oilTable = QTableWidget(self.oilFrame)

        self.oilTable.setColumnCount(6)
        self.oilTable.setHorizontalHeaderLabels(["اسم الصنف","كود","كمية الزيت الأصلية","كمية الزيت المتبقية","السعر","الحالة"])
        self.oilTable.setColumnWidth(0,83)
        self.oilTable.setColumnWidth(1,104)
        self.oilTable.setColumnWidth(2,104)
        self.oilTable.setColumnWidth(3,104)
        self.oilTable.setColumnWidth(4,83)

        self.contextMenuoilTable = QMenu(self.oilTable)
        self.contextMenuoilTable.setStyleSheet("background-color:grey")
        self.createButtonOilTable()

        self.oilTable.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self.oilTable.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.oilTable.setSelectionMode(QTableWidget.SelectionMode.SingleSelection)
        self.oilTable.customContextMenuRequested.connect(self.showMenuOilTable)

        #Start minMenu Buttons style
        self.mainMenuOilFrame = QFrame(self.oilFrame)
        self.mainMenuOilFrame.setStyleSheet("background-color:white;border:2px solid black")
        
        label = QLabel(self.mainMenuOilFrame,text="القائمة الرئيسيه")
        label.setStyleSheet("background-color:white;border-bottom:none;font: 14pt 'Arial';")
        label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        label.setGeometry(0,0,181,31)

        addOilButton = QPushButton(self.mainMenuOilFrame,text="اضافة زيت")
        addOilButton.setGeometry(0,50,181,31)
        addOilButton.setStyleSheet("QPushButton {background-color:white;border:2px solid black;font: 14pt 'Arial';} QPushButton:hover {background-color:#c8c8c8;}")
        addOilButton.clicked.connect(self.addOil)

        exportOilReportButton = QPushButton(self.mainMenuOilFrame,text="تصدير تقرير مستودع الزيت")
        exportOilReportButton.setGeometry(0,90,181,31)
        exportOilReportButton.setStyleSheet("QPushButton {background-color:white;border:2px solid black;font: 14pt 'Arial';} QPushButton:hover {background-color:#c8c8c8;}")
        exportOilReportButton.clicked.connect(self.exportOilReport)

        self.putOilFrameAndStuff()

        self.loadOil()
        self.oilFrame.show()
    def addSearchOilEntries(self):
        try:
            self.searchEntryOil.destroy()
            self.searchEntryOil.hide()
        except:
            pass
        if self.comboSearchOilBox.currentText() == "الكل":
            self.loadOil()
        else:
            self.searchEntryOil = QLineEdit(self.oilFrame)
            self.searchEntryOil.textChanged.connect(self.searchOilFun)
            self.searchEntryOil.setGeometry(300,30,150,20)
            self.searchEntryOil.show()
    def searchOilFun(self):
        if len(self.searchEntryOil.text())==0:
            self.loadOil()
        else:
            self.oilTable.setRowCount(0)
            tempThing = [] 
            cr.execute("SELECT code From oil")
            choices = cr.fetchall()
            posiple = []
            for o in choices:
                for n,i in enumerate(o):
                    try:
                        if o[n][:len(self.searchEntryOil.text())]==self.searchEntryOil.text():
                                if i not in posiple:
                                    posiple.append(i)
                    except:
                        pass

            for p in posiple:
                cr.execute("SELECT name,code,originalQuantity,remeaningQuantity,price,condition FROM oil WHERE code = ?", [p])
                for i in cr.fetchall():
                    tempThing.append(i)

            for row,i in enumerate(tempThing):
                self.oilTable.insertRow(self.oilTable.rowCount())
                for col,val in enumerate(i):
                    self.oilTable.setItem(row,col,QTableWidgetItem(str(val)))
            self.changeBackgroundColorsToOilTable()
    def exportOilReport(self):
        try:
            self.destroyFrame(self.exportOilReportFrame)
        except:
            pass

        arabic_locale = QLocale(QLocale.Language.Arabic, QLocale.Country.SaudiArabia)

        self.exportOilReportFrame = QFrame(parent=self.mainFrame)
        self.exportOilReportFrame.setGeometry((self.mainFrame.width()-210)//2,(self.mainFrame.height()-230)//2,210,230)
        self.exportOilReportFrame.setStyleSheet("background-color:white;border:2px solid black")

        closeButton = QPushButton(self.exportOilReportFrame)
        closeButton.setStyleSheet("QPushButton {border:2px solid black;font-size:8px;qproperty-icon: url('assests/close.png');qproperty-iconSize: 18px}QPushButton:hover {background-color:#c8c8c8;}")        
        closeButton.setGeometry(160,10,31,21)
        closeButton.clicked.connect(lambda x, frame=self.exportOilReportFrame:self.destroyFrame(frame))

        label = QLabel(parent=self.exportOilReportFrame,text="من تاريخ")
        label.setStyleSheet('font: 18pt "Arial";border:none')
        label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        label.move(50,50)

        self.fromDateExportOil = QDateEdit(parent=self.exportOilReportFrame)
        self.fromDateExportOil.setCalendarPopup(True)
        self.fromDateExportOil.setDisplayFormat("yyyy/MM/dd")

        self.todayButtonFromDateOil = QPushButton("اليوم",clicked=lambda:self.fromDateExportOil.calendarWidget().setSelectedDate(QDate().currentDate()))
        self.todayButtonFromDateOil.setStyleSheet("background-color:green;")

        self.fromDateExportOil.setLocale(arabic_locale)
        self.fromDateExportOil.setFont(QFont("Arial",12))
        self.fromDateExportOil.setStyleSheet("background-color:white;color:black")
        
        self.fromDateExportOil.calendarWidget().layout().addWidget(self.todayButtonFromDateOil)
        self.fromDateExportOil.calendarWidget().setSelectedDate(QDate().currentDate())
        self.fromDateExportOil.setGeometry(20,80,160,31)

        label = QLabel(parent=self.exportOilReportFrame,text="الى تاريخ")
        label.setStyleSheet('font: 18pt "Arial";border:none')
        label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        label.move(50,120)

        self.toDateExportOil = QDateEdit(parent=self.exportOilReportFrame)
        self.toDateExportOil.setCalendarPopup(True)
        self.toDateExportOil.setDisplayFormat("yyyy/MM/dd")

        self.todayButtonToDateOil = QPushButton("اليوم",clicked=lambda:self.toDateExportOil.calendarWidget().setSelectedDate(QDate().currentDate()))
        self.todayButtonToDateOil.setStyleSheet("background-color:green;")

        self.toDateExportOil.setLocale(arabic_locale)
        self.toDateExportOil.setFont(QFont("Arial",12))
        self.toDateExportOil.setStyleSheet("background-color:white;color:black")
        
        self.toDateExportOil.calendarWidget().layout().addWidget(self.todayButtonToDateOil)
        self.toDateExportOil.calendarWidget().setSelectedDate(QDate().currentDate())
        self.toDateExportOil.setGeometry(20,150,160,31)

        exportButton = QPushButton(parent=self.exportOilReportFrame,text="تصدير")
        exportButton.setGeometry(50,190,101,31)
        exportButton.setStyleSheet("QPushButton:hover {background-color:#c8c8c8;}")
        exportButton.clicked.connect(self.completeExportOilReport)


        self.exportOilReportFrame.show()
    def completeExportOilReport(self):
        filePath = QFileDialog.getExistingDirectory(self,"Select a Directory")
        if len(filePath) > 0:
            valuesToWrite = []
            oilIds = []
            totalPriceToWrite = 0
            fromDate = str(self.fromDateExportOil.text()).split("/")
            fromDateGoodFormat = date(int(fromDate[0]),int(fromDate[1]),int(fromDate[2]))

            toDate = str(self.toDateExportOil.text()).split("/")
            toDateGoodFormat = date(int(toDate[0]),int(toDate[1]),int(toDate[2]))

            #[itemPrice,itemName, installDate, bigTruckNumber]

            cr.execute("SELECT oilId,oilAddDate FROM oilDate")
            for i in cr.fetchall():
                date1 = str(i[1]).split("-")
                date2 = date(int(date1[0]),int(date1[1]),int(date1[2]))

                if date2>=fromDateGoodFormat and date2 <= toDateGoodFormat:
                    oilIds.append(i[0])

            for oilId in oilIds:
                cr.execute("SELECT name,code,price FROM oil WHERE code=?",[oilId])
                result = list(cr.fetchall()[0])
                cr.execute("SELECT oilAddDate FROM oilDate WHERE oilId=?",[oilId])
                result.append(cr.fetchone()[0])
                valuesToWrite.append([result[0],result[1],str(round(float(result[2]),2)),result[3]])
                totalPriceToWrite+=round(float(result[2]),2)

            doc = docx.Document()

            sections = doc.sections
            sections.page_height = 11.69
            sections.page_width = 8.27
            sections = sections[-1]
            sections.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE


            new_width,new_height = sections.page_height,sections.page_width
            sections.page_width = new_width
            sections.page_height = new_height

            sections = doc.sections

            for section in sections:
                section.top_margin = docx.shared.Cm(0.3)
                section.bottom_margin = docx.shared.Cm(0.3)
                section.left_margin = docx.shared.Cm(0.3)
                section.right_margin = docx.shared.Cm(0.3)

            benefits_table = doc.add_table(rows=1,cols=5)
            benefits_table.style = "Table Grid"
            hdr_Cells = benefits_table.rows[0].cells

            hdr_Cells[4].text = "م"
            hdr_Cells[3].text = "اسم الصنف"
            hdr_Cells[2].text = "الكود"
            hdr_Cells[1].text = "السعر"
            hdr_Cells[0].text = "تاريخ الشراء"

            for cell in hdr_Cells:
                self.set_arabic_format(cell)

            b = 0
            
            for i in valuesToWrite:
                b+=1   
                row_Cells = benefits_table.add_row().cells
                row_Cells[0].size = docx.shared.Pt(15)
                row_Cells[1].size = docx.shared.Pt(15)
                row_Cells[2].size = docx.shared.Pt(15)
                row_Cells[3].size = docx.shared.Pt(15)
                row_Cells[4].size = docx.shared.Pt(15)

                row_Cells[4].text = str(b)
                row_Cells[3].text = str(i[0])
                row_Cells[2].text = str(i[1])
                row_Cells[1].text = str(i[2])
                row_Cells[0].text = str(i[3])

                for cell in row_Cells:
                    self.set_arabic_format(cell)

            widths = (docx.shared.Inches(4), docx.shared.Inches(4), docx.shared.Inches(4), docx.shared.Inches(4),docx.shared.Inches(0.5))
            for row in benefits_table.rows:
                for idx, width in enumerate(widths):
                    row.cells[idx].width = width
            for row in benefits_table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size= docx.shared.Pt(17)
            para = doc.add_paragraph().add_run(f"اجمالي الصرف على مستودع الزيت:{round(totalPriceToWrite,2)}")
            para.font.name = "Arial"
            para.font.size = docx.shared.Pt(20)
            
            doc.save(f"{filePath}\تقرير مستودع الزيت.docx")
            with suppress_output():
                convert(f"{filePath}\تقرير مستودع الزيت.docx",f"{filePath}\تقرير مستودع الزيت.pdf")
            os.remove(f"{filePath}\تقرير مستودع الزيت.docx")

            message = QMessageBox(parent=self,text="تم التصدير بنجاح")
            message.setIcon(QMessageBox.Icon.Information)
            message.setWindowTitle("نجاح")
            message.exec()
    def putOilFrameAndStuff(self):
        
        if self.width() <= 995 or self.height() <= 680:
            self.oilFrame.setGeometry((self.width()-680)//2,(self.mainFrame.height()-610)//2,680,610)
            self.oilTable.setGeometry(10,60,441,541)
            self.mainMenuOilFrame.setGeometry(460,60,181,161)
            self.closeButtonOilFrame.setGeometry(630,10,41,31)
        else:
            self.oilFrame.setGeometry(30,30,self.mainFrame.width()-60,self.mainFrame.height()-60)
            self.oilTable.setGeometry(10,60,self.oilFrame.width() - 206,self.oilFrame.height()-100)
            self.mainMenuOilFrame.setGeometry(self.oilTable.width() + 20,60,181,161)
            self.closeButtonOilFrame.setGeometry(self.oilTable.width() + 20 + 140,10,41,31)

        deserve = self.oilTable.columnCount()
        for i in range(self.oilTable.columnCount()):
            self.oilTable.setColumnWidth(i, (self.oilTable.width() - 20) // deserve)

        self.oilFrame.show()
        self.oilTable.show()
        self.mainMenuOilFrame.show()
        self.closeButtonOilFrame.show()
    def showOilLifecycle(self):
        try:
            self.showOilLifecycycleFrame.deleteLater()
        except:
            pass

        self.oilIdShowLifecycle = self.oilTable.item(self.oilTable.selectedIndexes()[0].row(),1).text()
        self.oilNameShowLifecycle = self.oilTable.item(self.oilTable.selectedIndexes()[0].row(),0).text()

        self.showOilLifecycycleFrame = QFrame(self.oilFrame)
        self.showOilLifecycycleFrame.setGeometry((self.oilFrame.width()-500)//2,(self.oilFrame.height()-450)//2,500,450)
        # self.showTrucksNoteFrame.setStyleSheet("background-color:white;border:2px solid black")
        
        closeButton = QPushButton(self.showOilLifecycycleFrame)
        closeButton.setStyleSheet("QPushButton {border:2px solid black;font-size:12px;qproperty-icon: url('assests/close.png');qproperty-iconSize: 26px}QPushButton:hover {background-color:#c8c8c8;}")        
        closeButton.setGeometry(400,10,41,31)
        closeButton.clicked.connect(lambda x, frame=self.showOilLifecycycleFrame:self.destroyFrame(frame))

        self.showOilLifecycycleFrame.setStyleSheet("background-color:white")

        self.oilLifecycleTable = QTableWidget(self.showOilLifecycycleFrame)
        self.oilLifecycleTable.setGeometry(10,60,460,240)

        self.oilLifecycleTable.setColumnCount(2)
        self.oilLifecycleTable.setHorizontalHeaderLabels(["التاريخ","الحدث"])
        self.oilLifecycleTable.setColumnWidth(0,120)
        self.oilLifecycleTable.setColumnWidth(1,330)

        self.oilLifecycleTable.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self.oilLifecycleTable.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.oilLifecycleTable.setSelectionMode(QTableWidget.SelectionMode.SingleSelection)

        exportItemsLifecycleReportButton = QPushButton(self.showOilLifecycycleFrame,text="تصدير التقرير")
        exportItemsLifecycleReportButton.setGeometry(150,380,181,31)
        exportItemsLifecycleReportButton.setStyleSheet("QPushButton {background-color:white;border:2px solid black;font: 14pt 'Arial';} QPushButton:hover {background-color:#c8c8c8;}")
        exportItemsLifecycleReportButton.clicked.connect(self.exportOilLifecycleReport)



        #End minMenu Buttons style
        self.loadOilLifecycle()
        self.showOilLifecycycleFrame.show()
    def exportOilLifecycleReport(self):
        filePath = QFileDialog.getExistingDirectory(self,"Select a Directory")
        if len(filePath) > 0:
            cr.execute("SELECT date,action FROM itemsLifecycle WHERE itemId = ?",[self.oilIdShowLifecycle])
            values = cr.fetchall()
            doc = docx.Document()

            sections = doc.sections
            sections.page_height = 11.69
            sections.page_width = 8.27
            sections = sections[-1]
            sections.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE


            new_width,new_height = sections.page_height,sections.page_width
            sections.page_width = new_width
            sections.page_height = new_height

            sections = doc.sections

            for section in sections:
                section.top_margin = docx.shared.Cm(0.3)
                section.bottom_margin = docx.shared.Cm(0.3)
                section.left_margin = docx.shared.Cm(0.3)
                section.right_margin = docx.shared.Cm(0.3)

            para = doc.add_paragraph()

            paraRun = para.add_run(f"دورة حياة الصنف {self.oilNameShowLifecycle}")
            paraRun.font.name = "Arial"
            paraRun.font.size = docx.shared.Pt(20)

            para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

            benefits_table = doc.add_table(rows=1,cols=3)
            benefits_table.style = "Table Grid"
            hdr_Cells = benefits_table.rows[0].cells
            hdr_Cells[2].text = "م"
            hdr_Cells[1].text = "التاريخ"
            hdr_Cells[0].text = "الحدث"

            for cell in hdr_Cells:
                self.set_arabic_format(cell)

            b = 0
            
            for i in values:
                b+=1   
                row_Cells = benefits_table.add_row().cells
                row_Cells[0].size = docx.shared.Pt(15)
                row_Cells[1].size = docx.shared.Pt(15)
                row_Cells[2].size = docx.shared.Pt(15)

                row_Cells[2].text = str(b)
                row_Cells[1].text = str(i[0])
                row_Cells[0].text = str(i[1])

                for cell in row_Cells:
                    self.set_arabic_format(cell)

            widths = (docx.shared.Inches(6),docx.shared.Inches(2),docx.shared.Inches(0.5))
            for row in benefits_table.rows:
                for idx, width in enumerate(widths):
                    row.cells[idx].width = width
            for row in benefits_table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size= docx.shared.Pt(20)
                            
            benefits_table.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            
            doc.save(f"{filePath}\دورة حياة الصنف {self.oilNameShowLifecycle}.docx")
            with suppress_output():
                convert(f"{filePath}\دورة حياة الصنف {self.oilNameShowLifecycle}.docx",f"{filePath}\دورة حياة الصنف {self.oilNameShowLifecycle}.pdf")
            os.remove(f"{filePath}\دورة حياة الصنف {self.oilNameShowLifecycle}.docx")

            message = QMessageBox(parent=self,text="تم التصدير بنجاح")
            message.setIcon(QMessageBox.Icon.Information)
            message.setWindowTitle("نجاح")
            message.exec()
    def loadOilLifecycle(self):
        self.oilLifecycleTable.setRowCount(0)     
        
        cr.execute("SELECT date,action FROM itemsLifecycle WHERE itemId = ?",[self.oilIdShowLifecycle])
        tempThing = [] 
        for i in cr.fetchall():
            tempThing.append(i)
        for row,i in enumerate(tempThing):
            self.oilLifecycleTable.insertRow(self.oilLifecycleTable.rowCount())
            for col,val in enumerate(i):
                self.oilLifecycleTable.setItem(row,col,QTableWidgetItem(str(val))) 
    def loadItemToTrucksOil(self):
        self.trucksTableToChoiceOil.setRowCount(0)
        cr.execute("SELECT id, carNumber FROM trucks")
        
        tempThing = []
        for i in cr.fetchall():
            i = list(i)
            i.insert(1,"")
            i.append("شاحنة")
            tempThing.append(i)

        cr.execute("SELECT id, number FROM bigTrucks")
        for i in cr.fetchall():
            i = list(i)
            i.insert(1,"")
            i.append("تيدر")
            tempThing.append(i)

        for row in range(len(tempThing)):
            self.trucksTableToChoiceOil.insertRow(self.trucksTableToChoiceOil.rowCount())
            for col in range(self.trucksTableToChoiceOil.columnCount()):
                self.trucksTableToChoiceOil.setItem(row,col,QTableWidgetItem(str(tempThing[row][col])))
                if col==1:
                    button = QRadioButton()
                    button.clicked.connect(lambda ch,truckId=tempThing[row][0]:self.changeCurrentIdOil(truckId))
                    if tempThing[row][0] == self.currentIdTruckOil:
                        button.setChecked(True)
                    self.trucksTableToChoiceOil.setIndexWidget(self.trucksTableToChoiceOil.model().index(row,1),button)
    def changeCurrentIdOil(self,truckId):
        self.currentIdTruckOil = truckId
    def AddOilToTruck(self):
        try:
            self.destroyFrame(self.addOilToTrukFrame)
        except:
            pass
        self.currentIdTruckOil = None
        self.OilId = self.oilTable.item(self.oilTable.selectedIndexes()[0].row(),1).text()
        self.OilName = self.oilTable.item(self.oilTable.selectedIndexes()[0].row(),0).text()
        cr.execute("SELECT pricePerLiter FROM oil WHERE code=?",[self.OilId])
        self.OilPrice = cr.fetchone()[0]
        
        self.addOilToTrukFrame = QFrame(self.mainFrame)
        self.addOilToTrukFrame.setGeometry((self.mainFrame.width()-360)//2,(self.mainFrame.height()-550)//2,360,600)
        self.addOilToTrukFrame.setObjectName("addOilToTrukFrame")
        self.addOilToTrukFrame.setStyleSheet("QFrame#addOilToTrukFrame {background-color:white;border:2px solid black}")

        closeButton = QPushButton(self.addOilToTrukFrame)
        closeButton.setStyleSheet("QPushButton {border:2px solid black;font-size:12px;qproperty-icon: url('assests/close.png');qproperty-iconSize: 26px;background-color:#fdfdfd}QPushButton:hover {background-color:#c8c8c8;}")        
        closeButton.setGeometry(323,10,31,31)
        closeButton.clicked.connect(lambda x, frame=self.addOilToTrukFrame:self.destroyFrame(frame))

        label = QLabel(parent=self.addOilToTrukFrame, text="اختر الشاحنة")
        label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        label.setStyleSheet("font: 14pt 'Arial';background-color:white")
        label.setGeometry(108,40,151,31)

        self.comboBoxChoicesSearchOil = QComboBox(parent=self.addOilToTrukFrame)
        self.comboBoxChoicesSearchOil.setStyleSheet("background-color:#fdfdfd")
        self.comboBoxChoicesSearchOil.addItems(["الكل", "رقم الشاحنة"])
        self.comboBoxChoicesSearchOil.setGeometry(256,70,101,22)
        self.comboBoxChoicesSearchOil.currentIndexChanged.connect(self.addSearchEntryAddOilToTruck)


        self.trucksTableToChoiceOil = QTableWidget(parent=self.addOilToTrukFrame)
        self.trucksTableToChoiceOil.setStyleSheet("background-color:white")
        self.trucksTableToChoiceOil.setColumnCount(4)
        self.trucksTableToChoiceOil.setColumnHidden(0,True)
        self.trucksTableToChoiceOil.setColumnWidth(1,40)
        self.trucksTableToChoiceOil.setColumnWidth(2,140)
        self.trucksTableToChoiceOil.setColumnWidth(3,140)

        self.trucksTableToChoiceOil.setHorizontalHeaderLabels(["id", "", "رقم السيارة","نوعه"])
        self.trucksTableToChoiceOil.setGeometry(11,100,346,121)

        label = QLabel('تاريخ التركيب',self.addOilToTrukFrame)
        label.setStyleSheet("border:0px;background-color:#fdfdfd")
        label.move(140,230)
        label.setFont(QFont("Arial", 16))

        arabic_locale = QLocale(QLocale.Language.Arabic, QLocale.Country.SaudiArabia)

        self.installDateEntryOil = QDateEdit(self.addOilToTrukFrame)


        self.installDateEntryOil.setCalendarPopup(True)
        self.installDateEntryOil.setDisplayFormat("yyyy/MM/dd")

        todayButtonNoteDate = QPushButton("اليوم",clicked=lambda:self.installDateEntryOil.calendarWidget().setSelectedDate(QDate().currentDate()))
        todayButtonNoteDate.setStyleSheet("background-color:green;")

        self.installDateEntryOil.setLocale(arabic_locale)
        self.installDateEntryOil.setFont(QFont("Arial",12))
        self.installDateEntryOil.setStyleSheet("background-color:#fdfdfd;color:black;border:1px solid black")

        self.installDateEntryOil.calendarWidget().layout().addWidget(todayButtonNoteDate)
        self.installDateEntryOil.calendarWidget().setSelectedDate(QDate().currentDate())
        self.installDateEntryOil.setGeometry(80, 260, 210, 30)

        label = QLabel('مكان التركيب',self.addOilToTrukFrame)
        label.setStyleSheet("border:0px;background-color:#fdfdfd")
        label.move(140,290)
        label.setFont(QFont("Arial", 16))

        self.installPlaceOil = QLineEdit(self.addOilToTrukFrame)
        self.installPlaceOil.setStyleSheet("background-color:#fdfdfd")
        self.installPlaceOil.setFont(QFont("Arial", 18))
        self.installPlaceOil.setGeometry(80,320,210,30)

        label = QLabel('الكمية',self.addOilToTrukFrame)
        label.setStyleSheet("border:0px;background-color:#fdfdfd")
        label.move(165,360)
        label.setFont(QFont("Arial", 16))

        self.howManyEntryOil = QLineEdit(self.addOilToTrukFrame)
        self.howManyEntryOil.setStyleSheet("background-color:#fdfdfd")
        self.howManyEntryOil.setFont(QFont("Arial", 18))
        self.howManyEntryOil.setGeometry(80,400,210,30)
        self.howManyEntryOil.setValidator(QIntValidator(0,999999999))

        label = QLabel('ملاحظات',self.addOilToTrukFrame)
        label.setStyleSheet("border:0px;background-color:#fdfdfd")
        label.move(150,440)
        label.setFont(QFont("Arial", 16))

        self.notesTextOil = QTextEdit(self.addOilToTrukFrame)
        self.notesTextOil.setStyleSheet("background-color:#ffffff")
        self.notesTextOil.setGeometry(11,470,346,80)

        addItemToTruckButton = QPushButton(parent=self.addOilToTrukFrame, text="اضافة")
        addItemToTruckButton.setStyleSheet("background-color:#fdfdfd")
        addItemToTruckButton.clicked.connect(self.completeAddOilToTruck)
        addItemToTruckButton.setGeometry(108,560,131,31)

        self.loadItemToTrucksOil()
        self.addOilToTrukFrame.show()
    def addSearchEntryAddOilToTruck(self):
        if self.comboBoxChoicesSearchOil.currentText() == "الكل":
            try:
                self.oilSearchEntry.destroy()
                self.oilSearchEntry.close()
            except:
                pass
            self.loadItemToTrucksWareHouse()
        else:
            try:
                self.oilSearchEntry.destroy()
                self.oilSearchEntry.close()
            except:
                pass
            self.oilSearchEntry = QLineEdit(parent=self.addOilToTrukFrame)
            self.oilSearchEntry.textChanged.connect(self.searchTruckToAddOilToTruck)
            self.oilSearchEntry.setStyleSheet("background-color:white;")
            self.oilSearchEntry.setGeometry(120,70,121,20)
            self.oilSearchEntry.show()
    def searchTruckToAddOilToTruck(self):
        if len(self.oilSearchEntry.text())==0:
            self.loadItemToTrucksOil()
        else:
            self.trucksTableToChoiceOil.setRowCount(0)
            tempThing = [] 
            cr.execute("SELECT carNumber From trucks")
            choices = cr.fetchall()
            posiple = []
            for o in choices:
                for n,i in enumerate(o):
                    try:
                        if o[n][:len(self.oilSearchEntry.text())]==self.oilSearchEntry.text():
                                if i not in posiple:
                                    posiple.append(i)
                    except:
                        pass

            #bigTrucks search
            cr.execute("SELECT number From bigTrucks")
            choices = cr.fetchall()
            for o in choices:
                for n,i in enumerate(o):
                    try:
                        if o[n][:len(self.oilSearchEntry.text())]==self.oilSearchEntry.text():
                                if i not in posiple:
                                    posiple.append(i)
                    except:
                        pass

            for p in posiple:
                cr.execute("SELECT id, carNumber FROM trucks WHERE carNumber = ?", [p])
                for i in cr.fetchall():
                    i = list(i)
                    i.insert(1,"")
                    i.append("شاحنة")
                    tempThing.append(i)
                    
            #search bigh trcks
            for p in posiple:
                cr.execute("SELECT id, number FROM bigTrucks WHERE number = ?", [p])
                for i in cr.fetchall():
                    i = list(i)
                    i.insert(1,"")
                    i.append("تيدر")
                    tempThing.append(i)


            for row in range(len(tempThing)):
                self.trucksTableToChoiceOil.insertRow(self.trucksTableToChoiceOil.rowCount())
                for col in range(self.trucksTableToChoiceOil.columnCount()):
                    self.trucksTableToChoiceOil.setItem(row,col,QTableWidgetItem(str(tempThing[row][col])))
                    if col==1:
                        button = QRadioButton()
                        button.clicked.connect(lambda ch,truckId=tempThing[row][0]:self.changeCurrentIdOil(truckId))
                        if tempThing[row][0] == self.currentIdTruckOil:
                            button.setChecked(True)
                        self.trucksTableToChoiceOil.setIndexWidget(self.trucksTableToChoiceOil.model().index(row,1),button)
    def completeAddOilToTruck(self):
        self.carTypeOil = None
        for row in range(self.trucksTableToChoiceOil.rowCount()):
            if self.trucksTableToChoiceOil.cellWidget(row,1).isChecked():
                self.carTypeOil = self.trucksTableToChoiceOil.item(row,3).text()

        if (self.currentIdTruckOil != None and len(self.installDateEntryOil.text()) > 0 and len(self.installPlaceOil.text())) > 0 and self.carTypeOil!=None and len(self.howManyEntryOil.text()) > 0:

            if self.carTypeOil == 'شاحنة':
                cr.execute("SELECT truckId FROM trucksToOil WHERE oilCode = ? and truckId=?",[self.OilId, self.currentIdTruckOil])
            elif self.carTypeOil == 'تيدر':
                cr.execute("SELECT bigTruckId FROM bigTrucksToOil WHERE oilCode = ? and bigTruckId=?",[self.OilId, self.currentIdTruckOil])

            if len(cr.fetchall()) != 0:
                message = QMessageBox(parent=self,text="نوع الزيت هذا مركب على الشاحنة/التيدر بالفعل")
                message.setIcon(QMessageBox.Icon.Critical)
                message.setWindowTitle("فشل")
                message.exec()
                return
            
            cr.execute("SELECT remeaningQuantity FROM oil WHERE code=?",[self.OilId])
            if int(cr.fetchone()[0]) - int(self.howManyEntryOil.text()) >= 0:
                if self.carTypeOil == "شاحنة":
                    cr.execute("INSERT INTO trucksToOil (oilCode, truckId, itemPrice, itemName, installPlace,installDate,quantity,notes) VALUES (?,?,?,?,?,?,?,?)",(self.OilId, self.currentIdTruckOil, round(self.OilPrice * int(self.howManyEntryOil.text()),2), self.OilName, self.installPlaceOil.text(), self.installDateEntryOil.text(),self.howManyEntryOil.text(),self.notesTextOil.toPlainText()))
                elif self.carTypeOil == "تيدر":
                    cr.execute("INSERT INTO bigTrucksToOil (oilCode, bigTruckId, itemPrice, itemName, installPlace,installDate,quantity,notes) VALUES (?,?,?,?,?,?,?,?)",(self.OilId, self.currentIdTruckOil, round(self.OilPrice * int(self.howManyEntryOil.text()),2), self.OilName, self.installPlaceOil.text(), self.installDateEntryOil.text(),self.howManyEntryOil.text(),self.notesTextOil.toPlainText()))
                    
                cr.execute(f"UPDATE oil SET remeaningQuantity=remeaningQuantity-{int(self.howManyEntryOil.text())} WHERE code=?",[self.OilId])

                if self.carTypeOil == "شاحنة":
                    cr.execute("SELECT carNumber FROM trucks WHERE id=?",[self.currentIdTruckOil])
                elif self.carTypeOil == "تيدر":
                    cr.execute("SELECT number FROM bigTrucks WHERE id=?",[self.currentIdTruckOil])

                if self.carTypeOil == "شاحنة":
                    cr.execute("INSERT INTO itemsLifecycle (itemId,date,action) VALUES (?,?,?)",(self.OilId, str(date.today()).replace("-","/"), f"تم تعبئة كمية {self.howManyEntryOil.text()} من الزيت الى شاحنة رقم {cr.fetchone()[0]}"))
                elif self.carTypeOil == "تيدر":
                    cr.execute("INSERT INTO itemsLifecycle (itemId,date,action) VALUES (?,?,?)",(self.OilId, str(date.today()).replace("-","/"), f"تم تعبئة كمية {self.howManyEntryOil.text()} من الزيت الى تيدر رقم {cr.fetchone()[0]}"))
                
                con.commit()
                message = QMessageBox(parent=self,text="تمت الاضافة بنجاح")
                message.setIcon(QMessageBox.Icon.Information)
                message.setWindowTitle("نجاح")
                message.exec()
                self.loadOil()
                try:
                    for child in self.addOilToTrukFrame.children():
                        child.deleteLater()
                    self.addOilToTrukFrame.deleteLater()
                except:
                    pass
            else:
                message = QMessageBox(parent=self,text="كمية الزيت لاتتوفر بالمخزون")
                message.setIcon(QMessageBox.Icon.Critical)
                message.setWindowTitle("فشل")
                message.exec()  

        else:
            message = QMessageBox(parent=self,text="يرجى تعبئة جميع الحقول")
            message.setIcon(QMessageBox.Icon.Critical)
            message.setWindowTitle("فشل")
            message.exec()    
    def addOil(self):

        try:
            self.destroyFrame(self.addOilFrame)
        except:
            pass

        self.addOilFrame = QFrame(self.mainFrame)
        self.addOilFrame.setGeometry((self.mainFrame.width()-350)//2,(self.mainFrame.height()-506)//2,350,506)
        self.addOilFrame.setStyleSheet("background-color:white;border:2px solid black")

        closeButton = QPushButton(self.addOilFrame)
        closeButton.setStyleSheet("QPushButton {border:2px solid black;font-size:12px;qproperty-icon: url('assests/close.png');qproperty-iconSize: 26px}QPushButton:hover {background-color:#c8c8c8;}")        
        closeButton.setGeometry(300,10,41,31)
        closeButton.clicked.connect(lambda x, frame=self.addOilFrame:self.destroyFrame(frame))
        
        #Start scrolAria
        
        self.frame = QFrame()

        layout = QVBoxLayout()
        self.frame.setLayout(layout)


        label = QLabel("اسم الصنف")
        label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        label.setStyleSheet("border:none;font: 15pt 'Arial';")
        self.nameEntry = QLineEdit()

        layout.addWidget(label)
        layout.addWidget(self.nameEntry)


        label = QLabel("الكود")
        label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        label.setStyleSheet("border:none;font: 15pt 'Arial';")
        self.codeEntryOil = QLineEdit()

        layout.addWidget(label)
        layout.addWidget(self.codeEntryOil)


        label = QLabel("الكمية")
        label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        label.setStyleSheet("border:none;font: 15pt 'Arial';")

        self.quantityEntry = QLineEdit()
        self.quantityEntry.setValidator(QIntValidator(0,999999999))

        layout.addWidget(label)
        layout.addWidget(self.quantityEntry)

        label = QLabel("السعر")
        label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        label.setStyleSheet("border:none;font: 15pt 'Arial';")
        self.priceEntry = QLineEdit()
        self.priceEntry.setValidator(QIntValidator(0,999999999))

        layout.addWidget(label)
        layout.addWidget(self.priceEntry)

        addButton = QPushButton(text="اضافة")
        addButton.clicked.connect(self.completeAddOil)
        addButton.setStyleSheet("QPushButton:hover {background-color:#c8c8c8;}")        

        layout.addWidget(addButton)

        self.scroolAria = QScrollArea(self.addOilFrame)
        self.scroolAria.setWidget(self.frame)
        self.scroolAria.setStyleSheet("border:1px solid gray")
        self.scroolAria.move(20,50)
        self.scroolAria.resize(321,431)

        self.scroolAria.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOn)


        self.scroolAria.setWidgetResizable(True)


        #End scrolAria
        self.addOilFrame.show()
    def completeAddOil(self):
        if (len(self.nameEntry.text()) > 0 and len(self.quantityEntry.text()) > 0 and len(self.priceEntry.text())) > 0 and len(self.codeEntryOil.text()) > 0:
            try:
                pricePerLiter = round(int(self.priceEntry.text()) / int(self.quantityEntry.text()), 4)
                cr.execute("SELECT name FROM oil WHERE code=?",[self.codeEntryOil.text()])
                if cr.fetchone()==None:
                    cr.execute("INSERT INTO oil (name,code,originalQuantity,remeaningQuantity,price,pricePerLiter,condition) VALUES (?,?,?,?,?,?,?)", (self.nameEntry.text(),self.codeEntryOil.text(),self.quantityEntry.text(),self.quantityEntry.text(),self.priceEntry.text(),pricePerLiter,"جديد"))
                    cr.execute("INSERT INTO itemsLifecycle (itemId, date, action) values(?,?,?)",(self.codeEntryOil.text(),str(date.today()).replace("-","/"),f"تم اضافة {self.quantityEntry.text()} من الزيت الى النظام"))
                    cr.execute("INSERT INTO oilDate (oilId,oilAddDate) VALUES (?,?)",(self.codeEntryOil.text(),str(date.today())))
                    con.commit()
                    self.loadOil()
                    message = QMessageBox(parent=self,text="تمت الاضافة بنجاح")
                    message.setIcon(QMessageBox.Icon.Information)
                    message.setWindowTitle("نجاح")
                    message.exec()
                else:
                    message = QMessageBox(parent=self,text="يوجد بالفعل صنف بنفس رقم الكود")
                    message.setIcon(QMessageBox.Icon.Critical)
                    message.setWindowTitle("فشل")
                    message.exec()
            except:
                message = QMessageBox(parent=self,text="حدث خطأ ملاحظة كمية الزيت لايمكن ان تكون 0")
                message.setIcon(QMessageBox.Icon.Critical)
                message.setWindowTitle("فشل")
                message.exec()
        else:
            message = QMessageBox(parent=self,text="يرجى تعبئة جميع الحقول")
            message.setIcon(QMessageBox.Icon.Critical)
            message.setWindowTitle("فشل")
            message.exec()
    def createButtonOilTable(self):
        deleteButton = QAction(self.oilTable)
        deleteButton.setIcon(QIcon("assests/trash.png"))
        deleteButton.setText("حذف")
        deleteButton.setFont(QFont("Arial" , 12))
        deleteButton.triggered.connect(self.deleteOil)

        addOilToTruck = QAction(self.oilTable)
        addOilToTruck.setText("اضافة الى شاحنة")
        addOilToTruck.setFont(QFont("Arial" , 12))
        addOilToTruck.triggered.connect(self.AddOilToTruck)
    
        showOilLifecycleButton = QAction(self.oilTable)
        showOilLifecycleButton.setText("اظهار دورة حياة")
        showOilLifecycleButton.setFont(QFont("Arial" , 12))
        showOilLifecycleButton.triggered.connect(self.showOilLifecycle)

        changeStatusButton = QAction(self.oilTable)
        changeStatusButton.setText("تغيير الحالة")
        changeStatusButton.setFont(QFont("Arial" , 12))
        changeStatusButton.triggered.connect(self.changeOilStatus)

        showWhereUsed = QAction(self.oilTable)
        showWhereUsed.setText("اظهار اماكن الاستخدام")
        showWhereUsed.setFont(QFont("Arial" , 12))
        showWhereUsed.triggered.connect(self.showOilInTrucks)

        self.contextMenuoilTable.addAction(addOilToTruck)
        self.contextMenuoilTable.addAction(deleteButton)
        self.contextMenuoilTable.addAction(showWhereUsed)
        self.contextMenuoilTable.addAction(changeStatusButton)
        self.contextMenuoilTable.addAction(showOilLifecycleButton)
    def changeBackgroundColorsToOilTable(self):
        for row in range(self.oilTable.rowCount()):
            if self.oilTable.item(row,5).text() == 'تالف':
                for col in range(self.oilTable.columnCount()):
                    self.oilTable.item(row,col).setBackground(QColor(255,0,0))
            elif self.oilTable.item(row,5).text() == 'تم البيع':
                for col in range(self.oilTable.columnCount()):
                    self.oilTable.item(row,col).setBackground(QColor(0,255,0))
    def changeOilStatus(self):
        self.oilIdchangeStatus = self.oilTable.item(self.oilTable.selectedIndexes()[0].row(),1).text()
        try:
            self.changeOilStatusFrame.deleteLater()
        except:
            pass
        self.changeOilStatusFrame = QFrame(self.oilFrame)
        self.changeOilStatusFrame.setGeometry((self.oilFrame.width()-340)//2,(self.oilFrame.height()-160)//2,340,160)

        closeButton = QPushButton(self.changeOilStatusFrame)
        closeButton.setStyleSheet("QPushButton {border:2px solid black;font-size:12px;qproperty-icon: url('assests/close.png');qproperty-iconSize: 26px}QPushButton:hover {background-color:#c8c8c8;}")        
        closeButton.setGeometry(290,10,41,31)
        closeButton.clicked.connect(lambda x, frame=self.changeOilStatusFrame:self.destroyFrame(frame))

        self.changeOilStatusFrame.setStyleSheet("background-color:white;border: 2px solid black")

        label = QLabel('الحالة',self.changeOilStatusFrame)
        label.setStyleSheet("border:0px")
        label.move(150,20)
        label.setFont(QFont("Arial", 16))

        self.comboChangeOilStatusBox = QComboBox(self.changeOilStatusFrame)
        self.comboChangeOilStatusBox.setGeometry(60,50,200,30)
        self.comboChangeOilStatusBox.addItems(["جديد","تالف","تم البيع"])

        cr.execute("SELECT condition FROM oil WHERE code=?",[self.oilIdchangeStatus])
        self.comboChangeOilStatusBox.setCurrentText(cr.fetchone()[0])

        completeeditItemNoteButton = QPushButton("تغيير الحالة",self.changeOilStatusFrame)
        completeeditItemNoteButton.setStyleSheet("QPushButton {border:2px solid black;font:14pt 'Arial'}QPushButton:hover {background-color:#c8c8c8;}")        
        completeeditItemNoteButton.setGeometry(70,110,180,31)
        completeeditItemNoteButton.clicked.connect(self.completeEditOilStatus)

        #End minMenu Buttons style
        self.changeOilStatusFrame.show()
    def completeEditOilStatus(self):
        cr.execute("SELECT condition FROM oil WHERE code=?",[self.oilIdchangeStatus])
        oldOilStatus = cr.fetchone()[0]
        cr.execute("UPDATE oil set condition=? WHERE code=?",(self.comboChangeOilStatusBox.currentText(),self.oilIdchangeStatus))
        if self.comboChangeOilStatusBox.currentText()!="جديد":
            cr.execute("SELECT originalQuantity FROM oil WHERE code=?",[self.oilIdchangeStatus])
            cr.execute("UPDATE oil set remeaningQuantity=? WHERE code=?",(cr.fetchone()[0],self.oilIdchangeStatus))
            cr.execute("DELETE FROM bigTrucksToOil WHERE oilCode=?",[self.oilIdchangeStatus])
            cr.execute("DELETE FROM trucksToOil WHERE oilCode=?",[self.oilIdchangeStatus])
            cr.execute("INSERT INTO itemsLifecycle (itemId, date, action) values (?,?,?)",(self.oilIdchangeStatus, str(date.today()).replace("-","/"), f"تم افراغ الزيت من جميع الشاحنات والتيادر"))
            cr.execute("INSERT INTO itemsLifecycle (itemId, date, action) values (?,?,?)",(self.oilIdchangeStatus, str(date.today()).replace("-","/"), f"تم تغيير حالة المنتج من {oldOilStatus} الى {self.comboChangeOilStatusBox.currentText()}"))

        con.commit()
        message = QMessageBox(parent=self,text="تم التعديل بنجاح")
        message.setIcon(QMessageBox.Icon.Information)
        message.setWindowTitle("نجاح")
        message.exec()
        self.loadOil()
    def deleteOil(self):
        oilIdDelete = self.oilTable.item(self.oilTable.selectedIndexes()[0].row(),1).text()
        d = QMessageBox(parent=self,text=f"تأكيد حذف الزيت من النظام")
        d.setIcon(QMessageBox.Icon.Information)
        d.setWindowTitle("تأكيد")
        d.setStyleSheet("background-color:white")
        d.setStandardButtons(QMessageBox.StandardButton.Cancel|QMessageBox.StandardButton.Ok)
        important = d.exec()
        if important == QMessageBox.StandardButton.Ok:
            cr.execute("DELETE FROM oil WHERE code=?",[oilIdDelete])
            cr.execute("DELETE FROM itemsLifecycle WHERE itemId=?",[oilIdDelete])
            cr.execute("DELETE FROM bigTrucksToOil WHERE oilCode=?",[oilIdDelete])
            cr.execute("DELETE FROM trucksToOil WHERE oilCode=?",[oilIdDelete])
            cr.execute("DELETE FROM oilDate WHERE oilId=?",[oilIdDelete])
            con.commit()
            self.loadOil()
            d = QMessageBox(parent=self,text="تم الحذف بنجاح")
            d.setWindowTitle("نجاح")
            d.setIcon(QMessageBox.Icon.Information)
            d.setStyleSheet("background-color:white")
            d.exec()
    def showMenuOilTable(self, position):
        indexes = self.oilTable.selectedIndexes()
        for index in indexes:
            if self.oilTable.item(self.oilTable.selectedIndexes()[0].row(),5).text() == 'تالف' or self.oilTable.item(self.oilTable.selectedIndexes()[0].row(),5).text() == 'تم البيع':
                self.contextMenuoilTable.actions()[0].setVisible(False)
            else:
                self.contextMenuoilTable.actions()[0].setVisible(True)
            self.contextMenuoilTable.exec(self.oilTable.viewport().mapToGlobal(position))

    def destroyFrame(self,frame):
        for i in frame.children():
            i.deleteLater()

        frame.destroy()
        frame.deleteLater()
    def loadOil(self):
        self.oilTable.setRowCount(0)

        if self.comboOilStatusBox.currentText() != 'الكل':
            cr.execute("SELECT name,code,originalQuantity,remeaningQuantity,price,condition FROM oil WHERE condition=?",[self.comboOilStatusBox.currentText()])
        else:
            cr.execute("SELECT name,code,originalQuantity,remeaningQuantity,price,condition FROM oil")

        tempThing = [] 
        for i in cr.fetchall():
            tempThing.append(i)
        for row,i in enumerate(tempThing):
            self.oilTable.insertRow(self.oilTable.rowCount())
            for col,val in enumerate(i):
                self.oilTable.setItem(row,col,QTableWidgetItem(str(val)))
        self.changeBackgroundColorsToOilTable()
    def showOilInTrucks(self):
        try:
            self.showOilInTrucksFrame.deleteLater()
        except:
            pass

        self.showOilInTrucksFrame = QFrame(self.oilFrame)
        self.showOilInTrucksFrame.setGeometry((self.oilFrame.width()-370)//2,(self.oilFrame.height()-450)//2,370,450)
        # self.showTrucksNoteFrame.setStyleSheet("background-color:white;border:2px solid black")
        self.OilIdShowItemsInTrucks = self.oilTable.item(self.oilTable.selectedIndexes()[0].row(),1).text()

        closeButton = QPushButton(self.showOilInTrucksFrame)
        closeButton.setStyleSheet("QPushButton {border:2px solid black;font-size:12px;qproperty-icon: url('assests/close.png');qproperty-iconSize: 26px}QPushButton:hover {background-color:#c8c8c8;}")        
        closeButton.setGeometry(270,10,41,31)
        closeButton.clicked.connect(lambda x, frame=self.showOilInTrucksFrame:self.destroyFrame(frame))

        self.showOilInTrucksFrame.setStyleSheet("background-color:white")

        self.OilInTruksTable = QTableWidget(self.showOilInTrucksFrame)
        self.OilInTruksTable.setGeometry(10,60,300,240)

        self.OilInTruksTable.setColumnCount(3)
        self.OilInTruksTable.setHorizontalHeaderLabels(["رقم الشاحنة/التيدر","كمية الزيت","نوع المركبة"])
        self.OilInTruksTable.setColumnWidth(0,120)
        self.OilInTruksTable.setColumnWidth(1,100)
        self.OilInTruksTable.setColumnWidth(2,60)

        self.OilInTruksTable.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self.OilInTruksTable.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.OilInTruksTable.setSelectionMode(QTableWidget.SelectionMode.SingleSelection)

        #End minMenu Buttons style
        self.loadOilInTrucks()
        self.showOilInTrucksFrame.show()
    def loadOilInTrucks(self):
        self.OilInTruksTable.setRowCount(0)     
        trucksIds= []
        bigTrucksIds = []
        tempThing = []

        cr.execute("SELECT truckId FROM trucksToOil WHERE oilCode=?",[self.OilIdShowItemsInTrucks])
        for i in cr.fetchall():
            trucksIds.append(i[0])


        for truckId in trucksIds:
            cr.execute("SELECT carNumber FROM trucks WHERE id=?",[truckId])
            val = list(cr.fetchall()[0])
            cr.execute("SELECT quantity FROM trucksToOil WHERE truckId=?",[truckId])
            val.append(cr.fetchone()[0])
            val.append("شاحنة")
            tempThing.append(val)

        cr.execute("SELECT bigTruckId FROM bigTrucksToOil WHERE oilCode=?",[self.OilIdShowItemsInTrucks])
        for i in cr.fetchall():
            bigTrucksIds.append(i[0])


        for bigTruckId in bigTrucksIds:
            cr.execute("SELECT number FROM bigTrucks WHERE id=?",bigTruckId)
            val = list(cr.fetchall()[0])
            cr.execute("SELECT quantity FROM bigTrucksToOil WHERE bigTruckId=?",[bigTruckId])
            val.append(cr.fetchone()[0])
            val.append("تيدر")
            tempThing.append(val)

        for row,i in enumerate(tempThing):
            self.OilInTruksTable.insertRow(self.OilInTruksTable.rowCount())
            for col,val in enumerate(i):
                self.OilInTruksTable.setItem(row,col,QTableWidgetItem(str(val))) 